import tkinter as tk
from tkinter import messagebox
from package.room import Room
from package.apartment import Apartment
from package.building import Building
from docx import Document
from openpyxl import Workbook

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Помещения")

        self.create_widgets()

    def create_widgets(self):
        tk.Label(self.root, text="Длина (м):", bg= '#2d3250', foreground='#ffffff', font="Century 11 normal roman").grid(row=0, column=0)
        tk.Label(self.root, text="Ширина (м):", bg= '#2d3250', foreground='#ffffff', font="Century 11 normal roman").grid(row=1, column=0)
        tk.Label(self.root, text="Количество комнат (для квартиры):", bg= '#2d3250', foreground='#ffffff', font="Century 11 normal roman").grid(row=2, column=0)
        tk.Label(self.root, text="Количество этажей (для дома):", bg= '#2d3250', foreground='#ffffff', font="Century 11 normal roman").grid(row=3, column=0)

        self.dlina_entry = tk.Entry(self.root, bg='#424769', foreground='#ffffff', font="Century 11 normal roman", borderwidth=3, width=12)
        self.shirina_entry = tk.Entry(self.root, bg='#424769', foreground='#ffffff', font="Century 11 normal roman", borderwidth=3, width=12)
        self.komnaty_entry = tk.Entry(self.root, bg='#424769', foreground='#ffffff', font="Century 11 normal roman", borderwidth=3, width=12)
        self.etazhi_entry = tk.Entry(self.root, bg='#424769', foreground='#ffffff', font="Century 11 normal roman", borderwidth=3, width=12)

        self.dlina_entry.grid(row=0, column=1)
        self.shirina_entry.grid(row=1, column=1)
        self.komnaty_entry.grid(row=2, column=1)
        self.etazhi_entry.grid(row=3, column=1)

        self.result_label = tk.Label(self.root, text="Результат: ", bg='#2d3250', foreground='#ffffff',font= "Century 11 normal roman")
        self.result_label.grid(row=4, columnspan=4, column=0)

        self.room_button = tk.Button(self.root, text="Рассчитать (Комната)", command=self.calculate_room, bg='#676f9d', foreground='#ffffff', font="Century 10 normal roman", width=18)
        self.apartment_button = tk.Button(self.root, text="Рассчитать (Квартира)", command=self.calculate_apartment, bg='#676f9d', foreground='#ffffff', font="Century 10 normal roman", width=19)
        self.building_button = tk.Button(self.root, text="Рассчитать (Дом)", command=self.calculate_building, bg='#676f9d', foreground='#ffffff', font="Century 10 normal roman", width=18)
        self.save_button = tk.Button(self.root, text="Сохранить результат", command=self.save_report, bg='#676f9d', foreground='#ffffff', font="Century 10 normal roman", width=18)

        self.room_button.grid(row=5, column=0)
        self.apartment_button.grid(row=5, column=1)
        self.building_button.grid(row=6, columnspan=2)
        self.save_button.grid(row=7, columnspan=2)

    def calculate_room(self):
        try:
            dlina = float(self.dlina_entry.get())
            shirina = float(self.shirina_entry.get())
            komnata = Room(dlina, shirina)
            self.show_result(komnata)
        except ValueError:
            messagebox.showerror("Ошибка", "Введите корректные числа")

    def calculate_apartment(self):
        try:
            dlina = float(self.dlina_entry.get())
            shirina = float(self.shirina_entry.get())
            komnaty = int(self.komnaty_entry.get())
            kvartira = Apartment(dlina, shirina, komnaty)
            self.show_result(kvartira)
        except ValueError:
            messagebox.showerror("Ошибка", "Введите корректные числа")

    def calculate_building(self):
        try:
            dlina = float(self.dlina_entry.get())
            shirina = float(self.shirina_entry.get())
            etazhi = int(self.etazhi_entry.get())
            komnaty = int(self.komnaty_entry.get())
            zdanie = Building(dlina, shirina, etazhi, komnaty)
            self.show_result(zdanie)
        except ValueError:
            messagebox.showerror("Ошибка", "Введите корректные числа")

    def show_result(self, obj):
        ploshad = obj.calculate_area()
        teplomoshch = obj.calculate_heating_power()
        self.result_label.config(text=f"Площадь: {ploshad} кв.м, Мощность: {teplomoshch} Вт", bg='#424769', foreground='#ffffff', font="Meiryo 8 bold roman")
        self.result = f"Площадь: {ploshad} кв.м, Мощность: {teplomoshch} Вт"

    def save_report(self):
        try:
            # Сохранить в .docx
            self.save_as_docx(self.result)

            # Или сохранить в .xls
            self.save_as_xls(self.result)

            messagebox.showinfo("Успех", "Результаты сохранены!")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при сохранении: {str(e)}")

    def save_as_docx(self, result):
        doc = Document()
        doc.add_heading('Отчёт по расчёту площади и тепловой мощности', 0)
        doc.add_paragraph(result)
        doc.save("otchet.docx")

    def save_as_xls(self, result):
        wb = Workbook()
        ws = wb.active
        ws.title = "Отчёт"
        ws['A1'] = 'Площадь и тепловая мощность'
        ws['A2'] = result
        wb.save("otchet.xlsx")


if __name__ == "__main__":
    root = tk.Tk()
    root.resizable(False, False)
    root['bg'] = '#2d3250'
    photo = tk.PhotoImage(file='house.png')
    root.iconphoto(False, photo)
    app = App(root)
    root.mainloop()



